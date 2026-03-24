"""Parse materials catalogs from CSV bytes (via pandas), DOCX tables, or PDF tables."""

from __future__ import annotations

import io
import re
from pathlib import Path
from typing import Any

import pandas as pd

REQUIRED = frozenset({"material_id", "material_type", "size_value", "unit_cost_usd"})

# (canonical_column, normalized aliases for header matching)
_HEADER_RULES: list[tuple[str, frozenset[str]]] = [
    (
        "material_id",
        frozenset(
            {
                "material_id",
                "material id",
                "mat id",
                "material #",
                "material no",
                "material number",
                "sku",
                "item",
                "item id",
                "item code",
                "item #",
                "item no",
                "item number",
                "part #",
                "part no",
                "part number",
                "code",
                "id",
                "product id",
                "product code",
            }
        ),
    ),
    (
        "material_type",
        frozenset(
            {
                "material_type",
                "material type",
                "type",
                "category",
                "class",
                "mat type",
                "product type",
            }
        ),
    ),
    (
        "size_value",
        frozenset(
            {
                "size_value",
                "size value",
                "size",
                "length",
                "len",
                "dimension",
                "nominal",
            }
        ),
    ),
    (
        "unit_cost_usd",
        frozenset(
            {
                "unit_cost_usd",
                "unit cost",
                "unit cost usd",
                "cost",
                "price",
                "unit price",
                "unitprice",
                "each",
                "ea",
                "usd",
                "amount",
            }
        ),
    ),
    (
        "unit",
        frozenset(
            {
                "unit",
                "uom",
                "um",
                "units",
                "unit of measure",
                "u/m",
            }
        ),
    ),
    (
        "size_desc",
        frozenset(
            {
                "size_desc",
                "size desc",
                "description",
                "desc",
                "name",
                "product name",
            }
        ),
    ),
]


def _norm_header(s: str) -> str:
    t = (s or "").strip().lower()
    t = re.sub(r"\s+", " ", t)
    t = t.replace("$", "").strip()
    return t


def _canonical_for_header(cell: str) -> str | None:
    h = _norm_header(cell)
    if not h:
        return None
    for canon, aliases in _HEADER_RULES:
        if h == canon or h in aliases:
            return canon
        for a in aliases:
            if h == a or (len(a) > 2 and a in h):
                return canon
    return None


def _first_number(s: str) -> float | None:
    if s is None:
        return None
    m = re.search(r"-?\d+(?:\.\d+)?", str(s).replace(",", ""))
    if not m:
        return None
    try:
        return float(m.group(0))
    except ValueError:
        return None


def _coerce_float(val: Any) -> float:
    if val is None or (isinstance(val, float) and pd.isna(val)):
        return float("nan")
    if isinstance(val, (int, float)):
        return float(val)
    s = str(val).strip().replace(",", "")
    if not s:
        return float("nan")
    try:
        return float(s)
    except ValueError:
        n = _first_number(s)
        return float(n) if n is not None else float("nan")


def _table_to_rows(raw: list[list[str | None]]) -> list[list[str]]:
    out: list[list[str]] = []
    for row in raw:
        out.append([(c or "").strip() if c is not None else "" for c in row])
    return out


def _rows_to_dataframe(rows: list[list[str]]) -> tuple[pd.DataFrame | None, str | None]:
    if not rows or not any(any(c.strip() for c in r) for r in rows):
        return None, "No table rows found in document."

    best_header_idx = 0
    best_score = -1
    for i, row in enumerate(rows[: min(5, len(rows))]):
        mapped = [_canonical_for_header(c) for c in row]
        score = len({m for m in mapped if m is not None})
        if score > best_score:
            best_score = score
            best_header_idx = i

    header_row = rows[best_header_idx]
    col_to_idx: dict[str, int] = {}
    for j, cell in enumerate(header_row):
        canon = _canonical_for_header(cell)
        if canon and canon not in col_to_idx:
            col_to_idx[canon] = j

    if not REQUIRED.issubset(col_to_idx.keys()):
        missing = sorted(REQUIRED - set(col_to_idx.keys()))
        if best_score < 2 and len(header_row) >= 4:
            col_to_idx = {
                "material_id": 0,
                "material_type": 1,
                "size_value": 2,
                "unit_cost_usd": 3,
            }
            if len(header_row) > 4:
                col_to_idx["unit"] = 4
            data_start = best_header_idx + 1
        else:
            return (
                None,
                f"Could not map required columns {missing} from headers. "
                f"Expected columns like material_id, material_type, size_value, unit_cost_usd.",
            )
    else:
        data_start = best_header_idx + 1

    records: list[dict[str, Any]] = []
    for row in rows[data_start:]:
        if not any(c.strip() for c in row):
            continue
        rec: dict[str, Any] = {}
        for col, idx in col_to_idx.items():
            if idx < len(row):
                rec[col] = row[idx]
            else:
                rec[col] = ""
        mid = str(rec.get("material_id", "")).strip()
        if not mid or mid.lower() in ("total", "subtotal", "notes", "—", "-"):
            continue
        records.append(rec)

    if not records:
        return None, "No data rows found after the header row."

    df = pd.DataFrame(records)
    return _validate_materials_df(df)


def _validate_materials_df(df: pd.DataFrame) -> tuple[pd.DataFrame | None, str | None]:
    cols = set(df.columns)
    if not REQUIRED.issubset(cols):
        return None, f"Materials table must include columns: {sorted(REQUIRED)}"

    df = df.copy()
    df["material_id"] = df["material_id"].astype(str).str.strip()
    df["material_type"] = df["material_type"].astype(str).str.strip()
    df["size_value"] = df["size_value"].map(_coerce_float)
    df["unit_cost_usd"] = df["unit_cost_usd"].map(_coerce_float)

    if "unit" not in df.columns:
        df["unit"] = "lf"
    else:
        df["unit"] = df["unit"].astype(str).str.strip()
        df["unit"] = df["unit"].replace("", "lf")

    bad = df["material_id"].eq("") | df["size_value"].isna() | df["unit_cost_usd"].isna()
    if bad.all():
        return None, "Could not parse numeric size_value or unit_cost_usd for any row."
    df = df.loc[~bad].reset_index(drop=True)
    if df.empty:
        return None, "No valid material rows after parsing numbers."

    allowed = REQUIRED | {"unit", "size_desc", "volume_cuft"}
    keep = [c for c in df.columns if c in allowed]
    df = df[keep]

    return df, None


def _slug_id(text: str, idx: int) -> str:
    t = re.sub(r"[^a-zA-Z0-9]+", "_", text.strip())[:40].strip("_")
    return t if t else f"ITEM_{idx:04d}"


def _parse_docx_paragraphs(data: bytes, default_material_type: str) -> tuple[pd.DataFrame | None, str | None]:
    """One material description per paragraph (no Word table)."""
    try:
        from docx import Document
    except ImportError as e:  # pragma: no cover
        return None, f"python-docx is required for DOCX: {e}"

    doc = Document(io.BytesIO(data))
    records: list[dict[str, Any]] = []
    i = 0
    for para in doc.paragraphs:
        line = para.text.replace("\n", " ").strip()
        if not line or line.lower() in ("material list", "materials", "notes"):
            continue
        i += 1
        records.append(
            {
                "material_id": _slug_id(line, i),
                "material_type": default_material_type,
                "size_value": 1.0,
                "unit_cost_usd": 0.0,
                "unit": "lf",
                "size_desc": line[:500],
            }
        )
    if not records:
        return None, "No material lines found in document."
    df = pd.DataFrame(records)
    return _validate_materials_df(df)


def parse_docx_tables(
    data: bytes,
    *,
    default_material_type: str = "misc",
) -> tuple[pd.DataFrame | None, str | None]:
    try:
        from docx import Document
    except ImportError as e:  # pragma: no cover
        return None, f"python-docx is required for DOCX: {e}"

    doc = Document(io.BytesIO(data))
    all_rows: list[list[str]] = []
    for table in doc.tables:
        for row in table.rows:
            all_rows.append([cell.text.replace("\n", " ").strip() for cell in row.cells])
    if not all_rows:
        return _parse_docx_paragraphs(data, default_material_type)
    return _rows_to_dataframe(all_rows)


def parse_pdf_tables(
    data: bytes,
    *,
    default_material_type: str = "misc",
) -> tuple[pd.DataFrame | None, str | None]:
    try:
        import pdfplumber
    except ImportError as e:  # pragma: no cover
        return None, f"pdfplumber is required for PDF: {e}"

    all_rows: list[list[str]] = []
    with pdfplumber.open(io.BytesIO(data)) as pdf:
        for page in pdf.pages:
            tables = page.extract_tables() or []
            for tbl in tables:
                all_rows.extend(_table_to_rows(tbl))
    if not all_rows:
        lines: list[str] = []
        with pdfplumber.open(io.BytesIO(data)) as pdf:
            for page in pdf.pages:
                text = (page.extract_text() or "").splitlines()
                lines.extend(t.strip() for t in text if t.strip())
        if not lines:
            return None, "No tables or text lines found in PDF."
        records = []
        for idx, line in enumerate(lines, start=1):
            if line.lower() in ("material list", "materials", "notes"):
                continue
            records.append(
                {
                    "material_id": _slug_id(line, idx),
                    "material_type": default_material_type,
                    "size_value": 1.0,
                    "unit_cost_usd": 0.0,
                    "unit": "lf",
                    "size_desc": line[:500],
                }
            )
        if not records:
            return None, "No material lines found in PDF."
        return _validate_materials_df(pd.DataFrame(records))
    return _rows_to_dataframe(all_rows)


def parse_catalog_bytes(
    filename: str,
    data: bytes,
    *,
    default_material_type: str = "misc",
) -> tuple[pd.DataFrame | None, str | None]:
    """
    Parse uploaded catalog bytes. Returns (df, error). On success error is None.
    """
    suf = Path(filename or "").suffix.lower()
    if suf == ".csv":
        try:
            df = pd.read_csv(io.BytesIO(data))
        except Exception as e:  # noqa: BLE001
            return None, str(e)
        return _validate_materials_df(df)
    if suf == ".docx":
        return parse_docx_tables(data, default_material_type=default_material_type)
    if suf == ".pdf":
        return parse_pdf_tables(data, default_material_type=default_material_type)
    return None, f"Unsupported materials file type: {suf or '(no extension)'}"
